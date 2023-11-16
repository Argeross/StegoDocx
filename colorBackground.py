import os
import hashlib
from docx import Document
from docx.shared import RGBColor, Pt

def get_checksum(text): # SHA-256 * len(txt)
    checksum = hashlib.sha256()
    checksum.update(text.encode())
    checksum = hex(int(checksum.hexdigest(), 16)*len(text))[2:]
    return checksum


def encode_to_bg(template_file, secret_txt, bg_color):
    # Get content of template file
    template_doc = Document(template_file)
    template_txt = "\n".join([paragraph.text for paragraph in template_doc.paragraphs])
    
    # Calculate checksum
    checksum = get_checksum(secret_txt)
    
    stego_txt = secret_txt + ":" + checksum
    
    # Create stego doc
    stego_doc = Document()
    stego_doc.add_paragraph(template_txt)
    stego_doc.add_paragraph(stego_txt)
    stego_doc.paragraphs[-1].runs[0].font.color.rgb = bg_color
    stego_doc.paragraphs[-1].runs[0].font.size = Pt(1)

    # Disable spell-checking in Word
    DOCX = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    element = stego_doc.settings.element.find(DOCX + 'proofState')
    element.attrib[DOCX + 'grammar'] = 'dirty'
    element.attrib[DOCX + 'spelling'] = 'dirty'
    stego_doc.core_properties.identifier = str(bg_color)
    
    # Saving file to the same directory as a template file
    full_path, ext = os.path.splitext(template_file)
    dir = os.path.dirname(full_path)
    stego_filename = os.path.splitext(os.path.basename(template_file))[0] + "_stego" + ext
    stego_full_path = os.path.join(dir, stego_filename)
    stego_doc.save(stego_full_path)
    
    #print(stego_full_path)
    return stego_full_path


def integrity_check(stego_hash, stego_text):
    return stego_hash == get_checksum(stego_text)

def decode_from_bg(stego_file):
    stego_doc = Document(stego_file)

    bg_color = RGBColor.from_string(stego_doc.core_properties.identifier)
    for paragraph in stego_doc.paragraphs:
        for run in paragraph.runs:
            if run.font.color.rgb == bg_color:
                stego_content = run.text
                break
    
    # try:
    stego_text, stego_hash = stego_content.split(":")
    # except:
        # return "Couldn't decode the message - error occured"

    if not integrity_check(stego_hash, stego_text):
        return "The checksum of a stegotext is invalid, probably the content was modified"
    
    return stego_text


if __name__ == "__main__":
    encode_to_bg("test.docx", "Testujemy", RGBColor(0, 0, 0))
    print(decode_from_bg("test_stego.docx"))