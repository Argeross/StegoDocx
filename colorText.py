import os
import hashlib
from docx import Document
from docx.text.run import Run
from docx.text.paragraph import Paragraph
from docx.shared import RGBColor
import string

LOWERCASE_LETTERS = {letter: RGBColor(0, 0, 0+i+1) for (i, letter) in enumerate(string.ascii_lowercase)}
UPPERCASE_LETTERS = {letter: RGBColor(0, 1, 0+i+1) for (i, letter) in enumerate(string.ascii_uppercase)}
NUMBERS = {str(number): RGBColor(0, 0+i+1, 0) for (i, number) in enumerate(range(0,10))}
SPECIAL_CHARS = {s_char: RGBColor(0+i+1, 0, 0) for (i, s_char) in enumerate(string.punctuation+" ")}

# Dict with characters mapped to RGB color
CHARS = {**LOWERCASE_LETTERS, **UPPERCASE_LETTERS, **NUMBERS, **SPECIAL_CHARS}
# End of the stegotext
END = RGBColor(1, 1, 1)

def get_checksum(text): # SHA-256 * len(txt)
    checksum = hashlib.sha256()
    checksum.update(text.encode())
    checksum = hex(int(checksum.hexdigest(), 16)*len(text))[2:]
    return checksum

def encode_to_color(template_file, secret_txt):
    # Get content of template file
    template_doc = Document(template_file)
    template_txt = "\n".join([paragraph.text for paragraph in template_doc.paragraphs])
    print(template_txt)
    
    # Calculate checksum
    checksum = get_checksum(secret_txt)
        
    if len(template_txt.replace("\n", "")) < len(secret_txt):
        print("Too short template file or too long message.")
        return "Too short template file or too long message."
    
    # Create stego doc
    stego_txt = secret_txt + ":" + checksum
    stego_doc = Document()

    # Color text
    i = 0
    prgph_num = 0
    complete = 0
    for p in template_doc.paragraphs:
        paragraph = stego_doc.add_paragraph()
        for run in p.runs:
            for idx, ltr in enumerate(run.text):
                if i >= len(stego_txt):

                    rest = paragraph.add_run()
                    rest.text = run.text[idx:]
                    rest.font.color.rgb = END # Color for the end of the message
                    complete = 1
                    break
                
                letter = paragraph.add_run()
                letter.text = ltr
                letter.font.color.rgb = CHARS[stego_txt[i]]
                i+=1
            
            if complete:
                break

        prgph_num+=1

        # Add the rest of template file
        if complete:
            for tparagraph in template_doc.paragraphs[prgph_num:]:
                stego_doc.add_paragraph(tparagraph.text)
            break

    # Saving file to the same directory as a template file
    full_path, ext = os.path.splitext(template_file)
    dir = os.path.dirname(full_path)
    stego_filename = os.path.splitext(os.path.basename(template_file))[0] + "_stego" + ext
    stego_full_path = os.path.join(dir, stego_filename)
    stego_doc.save(stego_full_path)
    
    print(stego_full_path)
    return stego_full_path

def integrity_check(stego_hash, stego_text):
    return stego_hash == get_checksum(stego_text)

def decode_from_color(stego_file):
    stego_doc = Document(stego_file)

    stego_content = ''
    complete = 0
    for paragraph in stego_doc.paragraphs:
        for run in paragraph.runs:
            if run.font.color.rgb != END:
                stego_content += list(CHARS.keys())[list(CHARS.values()).index(run.font.color.rgb)] # Returns characters associated with the color
            else:
                complete = 1
                break
        
        if complete:
            break
    
    try:
        stego_text, stego_hash = stego_content.split(":")
    except:
        return "Couldn't decode the message - error occured"

    if not integrity_check(stego_hash, stego_text):
        return "The checksum of a stegotext is invalid, probably the content was modified"
    
    return stego_text    


if __name__ == "__main__":
    encode_to_color("template.docx", "Testujemy 67&!?#")
    print(decode_from_color("template_stego.docx"))