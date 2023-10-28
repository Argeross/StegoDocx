import os
import hashlib
from docx import Document
from docx.shared import RGBColor



def get_checksum(text): # SHA-256 * len(txt)
    checksum = hashlib.sha256()
    checksum.update(text.encode())
    checksum = hex(int(checksum.hexdigest(), 16)*len(text))[2:]
    return checksum


def encode_in_spaces(template_file, secret_txt):
    # Get content of template file
    template_doc = Document(template_file)
    template_txt = "\n".join([paragraph.text for paragraph in template_doc.paragraphs]).strip()
    # print(template_txt)
    # Calculate checksum
    checksum = get_checksum(secret_txt)
    
    stego_txt = secret_txt + ":" + checksum
    txt_bin = ''.join(format(ord(c), '08b') for c in stego_txt)
    print(txt_bin)

    template_txt = [*template_txt]
    # print(template_txt)
    if template_txt.count(' ') < len(txt_bin):
        print("Too short template file or too long message.")
        return "Too short template file or too long message."

    encoded_txt = ""
    for i in range(len(txt_bin)):
        # print(template_txt[char_idx:])
        for j in range(len(template_txt)):
            # print(char)
            encoded_txt += template_txt[j]
            if template_txt[j] == ' ':
                if txt_bin[i] == '0': # one space == 1, two spaces == 0
                    encoded_txt += ' '
                    template_txt = template_txt[j+1:]
                    break
                template_txt = template_txt[j+1:]
                break
    encoded_txt += '.'
    # print(encoded_txt)
        

    # Create stego doc
    stego_doc = Document()
    stego_doc.add_paragraph(encoded_txt)

    # Disable spell-checking in Word
    DOCX = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    element = stego_doc.settings.element.find(DOCX + 'proofState')
    element.attrib[DOCX + 'grammar'] = 'dirty'
    element.attrib[DOCX + 'spelling'] = 'dirty'
    
    # Saving file to the same directory as a template file
    full_path, ext = os.path.splitext(template_file)
    dir = os.path.dirname(full_path)
    stego_filename = os.path.splitext(os.path.basename(template_file))[0] + "_stego" + ext
    stego_full_path = os.path.join(dir, stego_filename)
    stego_doc.save(stego_full_path)
    
    print(stego_full_path)
    return stego_full_path


def integrity_check(stego_hash, stego_text):
    return stego_hash == get_checksum(stego_text)

def decode_from_spaces(stego_file):
    stego_doc = Document(stego_file)

    stego_txt = "\n".join([paragraph.text for paragraph in stego_doc.paragraphs])
    print(stego_txt)
    decoded_bin = ''
    for i in range(len(stego_txt)):
        if stego_txt[i] == ' ':
            if stego_txt[i+1] == ' ':
                decoded_bin += '0'
            elif stego_txt[i-1] == ' ':
                pass
            else:
                decoded_bin += '1'
    print(decoded_bin)
    decoded_text = ''
    for i in range(0, len(decoded_bin), 8):
        byte = decoded_bin[i:i + 8]
        decoded_text += chr(int(byte, 2))
    

    print(decoded_text)
    try:
        message, stego_hash = decoded_text.split(":")
    except:
        return "Couldn't decode the message - error occured"


    if not integrity_check(stego_hash, message):
        return "The checksum of a stegotext is invalid, probably the content was modified"
    
    return message


if __name__ == "__main__":
    encode_in_spaces("template.docx", "Testujemy")
    # print(decode_from_spaces("template_stego.docx"))